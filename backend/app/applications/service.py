import os
import mimetypes

from pathlib import Path
import re
from docx import Document
from docx.shared import Pt
from datetime import datetime
from loguru import logger
from fastapi.responses import FileResponse

from app.applications.dao import ApplicationsDAO
from app.applications.models import ApplicationStatus, Applications
from app.notification.service import make_notification
from app.service.base import convert_to_pdf

async def create_applications(
    tg_id: int,
    fio: str,
    phone: str,
    email: str,
    cadastral_number: str,
    problem: str,
    address: str,
):
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    safe_filename = f"{timestamp}.docx"
    upload_dir = "doc/applications"

    file_path = os.path.join(upload_dir, safe_filename)

    try:
        applicant = f"{fio.split(' ')[0]} {fio.split(' ')[1][0]}. {fio.split(' ')[2][0]}."
    except Exception as e:
        raise {"detail": "Неверный формат имени"}

    # Создаем шаблон
    applicant_statement(
        output_path=file_path,
        applicant=applicant,
        email=email,
        phone_number=phone,
        problem=problem,
        address=address,
    )

    await convert_to_pdf(input_path=file_path)

    street = (re.search(r"ул\.\s[^,]+", address)).group()
    application_data: Applications = await ApplicationsDAO.add(
        tg_id=tg_id,
        fio=fio,
        phone=phone,
        email=email,
        cadastral_number=cadastral_number,
        street=street,
        address=address,
        file_url=file_path,
    )

    await make_notification(
        tg_id=tg_id,
        text=f"Ваше заявление №{application_data.id} принято",
    )

    return {"detail": "Заявление успешно создано", "application": application_data}


async def delete_applications(id: int):
    application_data = await ApplicationsDAO.find_by_id(id)
    if not application_data:
        raise {"detail": "Заявка не найдена"}
    if application_data.status != ApplicationStatus.ACCEPTED:
        raise {"detail": "Невозможно удалить заявление"}
    application_data = await ApplicationsDAO.delete(id)
    try:
        os.remove(application_data.file_url)
        os.remove(f"{(application_data.file_url).split(".")[1]}.pdf")
    except Exception as e:
        logger.error(f"Ошибка удаления: {e}")
    return {"detail": "Заявление успешно удалено", "application": application_data}


async def all_applications():
    application_data = await ApplicationsDAO.all()
    return {"count": len(application_data), "applications": application_data}


async def download_applications(id: int):
    application_data = await ApplicationsDAO.find_by_id(id)
    if not application_data:
        return {"detail": "Файл не найден"}
    return application_data.file_url


async def view_applications(id: int):
    application_data = await ApplicationsDAO.find_by_id(id)
    if not application_data:
        return {"detail": "Файл не найден"}
    file_path = f"{(application_data.file_url).split('.')[0]}.pdf"
    ext = os.path.splitext(file_path)[1].lower()
    try:
        media_type, _ = mimetypes.guess_type(file_path)
        if not media_type:
            media_type = "application/octet-stream"
    except Exception as e:
        raise {"detail": "Ошибка загрузки файла"}
    return FileResponse(path=file_path, media_type=media_type)


async def search_applications(text: str):
    applications_data = await ApplicationsDAO.search(text)
    return {"count": len(applications_data), "applications": applications_data}


async def filter_applications(
    street: str | None, date_from: datetime, date_to: datetime, is_departure: bool
):
    applications_data = await ApplicationsDAO.filter(
        street, date_from, date_to, is_departure
    )
    return {"count": len(applications_data), "applications": applications_data}


async def update_departure(applications_id: str, departure_date: datetime):
    if applications_id.count("-") == 1 and applications_id.count(",") <= 0:
        applications_id_list = list(
            range(
                int((applications_id.split("-"))[0]),
                (int((applications_id.split("-"))[-1])) + 1,
            )
        )
    elif applications_id.count(",") > 0 and applications_id.count("-") == 0:
        applications_id_list = [int(j) for j in applications_id.split(",")]
    else:
        try:
            applications_id_list = [int(applications_id)]
        except Exception as e:
            raise {"detail": "Неверный формат ввода"}

    applications_data = await ApplicationsDAO.departure(
        applications_id_list, departure_date
    )

    for application in applications_data:
        logger.debug(f"{application['tg_id']}, {application['id']}")
        await make_notification(
            tg_id=application["tg_id"],
            text=f"Для заявки №{application['id']} назначен выезд на {departure_date.strftime('%d.%m.%Y')}",
        )
    if len(applications_data) < 1:
        return {"detail": "Выезды не найдены или уже назначены"}

    return {"detail": "Выезд успешно назначен"}


def applicant_statement(
    output_path,
    applicant: str,
    email: str,
    phone_number: str,
    problem: str,
    address: str,
):
    try:
        template_path = "doc/templates/applications_templates.docx"
        # Преобразуем строковые пути в объекты Path
        template_path = Path(template_path)
        output_path = Path(output_path)

        # Проверяем существование шаблона
        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")

        doc = Document(template_path)

        # автоматически определяем сегодняшнюю дату
        date = datetime.today().strftime("%d.%m.%Y")

        replacements = {
            "{DATE}": date,
            "{APPLICANT}": applicant,
            "{EMAIL}": email,
            "{PHONE_NUMBER}": phone_number,
            "{PROBLEM}": problem,
            "{ADDRESS}": address,
        }

        def replace_in_paragraph(paragraph):
            text = paragraph.text
            for key, value in replacements.items():
                if key in text:
                    text = text.replace(key, value)
            if text != paragraph.text:
                paragraph.clear()
                run = paragraph.add_run(text)
                run.font.underline = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        # обработка абзацев
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)

        # обработка таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

        # ==== БЛОК С ДАТОЙ И ПОДПИСЬЮ ====
        table = doc.add_table(rows=2, cols=3)
        table.autofit = True

        row = table.rows[0].cells
        r0 = row[0].paragraphs[0].add_run(date)
        r0.font.name = "Times New Roman"
        r0.font.size = Pt(12)

        r1 = row[1].paragraphs[0].add_run("______________")
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(12)

        r2 = row[2].paragraphs[0].add_run(applicant)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(12)

        # пояснения
        row = table.rows[1].cells
        r1 = row[1].paragraphs[0].add_run("   подпись")
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(12)

        # Создаем папку для выходного файла, если её нет
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    except FileNotFoundError as e:
        logger.error(f"Ошибка: {e}")
    except Exception as e:
        logger.error(f"Произошла ошибка при обработке документа: {e}")
