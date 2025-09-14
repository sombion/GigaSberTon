import os
import mimetypes

from docx import Document

from datetime import datetime

from fastapi.responses import FileResponse
from app.applications.dao import ApplicationsDAO
from app.applications.models import ApplicationStatus, Applications
from app.auth.dao import UsersDAO
from app.auth.models import Users
from app.conclusion.dao import ConclusionDAO
from app.conclusion.models import Conclusion
from app.notification.dao import NotificationDAO
from app.notification.service import make_notification
from app.service.base import convert_to_pdf
from app.signature.dao import SignatureDAO


async def create_conclusions(
    applications_id: int,
    date: datetime,
    chairman_id: int,
    members_id: list,
    justification: str,
    documents: str,
    conclusion: str,
):
    chairman_data: Users = await UsersDAO.find_by_id(chairman_id)
    if not chairman_data:
        raise {"detail": "Пользователь не найден"}
    chairman_fio = chairman_data.fio
    try:
        chairman = f"{chairman_fio.split(' ')[0]} {chairman_fio.split(' ')[1][0]}. {chairman_fio.split(' ')[2][0]}."
    except Exception as e:
        raise {"detail": "Неверный формат имени"}

    members = []
    for member_id in members_id:
        member_data: Users = await UsersDAO.find_by_id(int(member_id))
        if not member_data:
            raise {"detail": "Пользователь не найден"}
        member_fio = member_data.fio
        try:
            member = f"{member_fio.split(' ')[0]} {member_fio.split(' ')[1][0]}. {member_fio.split(' ')[2][0]}."
        except Exception as e:
            raise {"detail": "Неверный формат имени"}
        members.append(member)

    applications_data: Applications = await ApplicationsDAO.find_by_id(applications_id)
    if not applications_data:
        raise {"detail": "Заявление не найдено"}

    if await ConclusionDAO.find_one_or_none(applications_id=applications_id):
        raise {"На данное заявление уже составленно заключение"}

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    safe_filename = f"{timestamp}.docx"
    upload_dir = "doc/conclusion"
    file_path = os.path.join(upload_dir, safe_filename)

    await fill_statement(
        file_path,
        str(datetime.now().strftime("%d.%m.%Y")),
        applications_data.address,
        "Администрацией Центрального района г. Воронежа, решение №123 от 01.09.2025",
        chairman,
        members,
        f"{applications_data.fio}, собственник квартиры",
        documents,
        conclusion,
        justification,
    )

    await convert_to_pdf(input_path=file_path)

    conclusion_data: Conclusion = await ConclusionDAO.add(
        applications_id, date, file_path
    )
    conclusion_id = conclusion_data.id

    # Изменение статуса
    await ApplicationsDAO.update(applications_id, ApplicationStatus.COMMISSION_REVIEW)

    # Уведомление для пользователя
    await NotificationDAO.add(
        user_id=chairman_id,
        text=f"Заявление №{conclusion_id} - вас назначили председателем комиссии",
    )
    await SignatureDAO.add(chairman_id, conclusion_id)
    for member_id in members_id:
        await NotificationDAO.add(
            user_id=int(member_id),
            text=f"Заявление №{conclusion_id} - вас назначили членом комиссии",
        )
        await SignatureDAO.add(int(member_id), conclusion_id)
    await make_notification(
        tg_id=applications_data.tg_id,
        text=f"На ваше заявление №{conclusion_id} создано заключение комиссии",
    )

    return {"detail": "Заявление комиссии успешно создано"}


async def edit_conclusions(): ...


async def all_conclusions():
    conclusion_data = await ConclusionDAO.all()
    return {"count": len(conclusion_data), "conclusions": conclusion_data}


async def search_conclusions(text: str):
    conclusion_data = await ConclusionDAO.search(text)
    return {"count": len(conclusion_data), "conclusions": conclusion_data}


async def filter_conclusions(
    street: str | None,
    date_from: datetime | None,
    date_to: datetime | None,
    signed: bool | None,
    user_id: int,
):
    conclusion_data = await ConclusionDAO.filter(street, date_from, date_to, signed, user_id)
    return {"count": len(conclusion_data), "conclusions": conclusion_data}


async def view_conclusions(id: int):
    conclusion_data = await ConclusionDAO.find_by_id(id)
    if not conclusion_data:
        return {"detail": "Файл не найден"}
    file_path = f"{(conclusion_data.file_url).split('.')[0]}.pdf"
    ext = os.path.splitext(file_path)[1].lower()
    try:
        media_type, _ = mimetypes.guess_type(file_path)
        if not media_type:
            media_type = "application/octet-stream"
    except Exception as e:
        raise {"detail": "Ошибка загрузки файла"}
    return FileResponse(path=file_path, media_type=media_type)


async def download_conclusions(id: int):
    conclusion_data = await ConclusionDAO.find_by_id(id)
    if not conclusion_data:
        return {"detail": "Файл не найден"}
    return conclusion_data.file_url


async def fill_statement(
    output_path,
    date: str,
    address: str,
    commission_info: str,
    chairman: str,
    members: list,
    owner: str,
    documents: str,
    conclusion: str,
    justification: str,
):
    doc = Document("doc/templates/statement_templates.docx")

    replacements = {
        "{DATE}": date,
        "{ADDRESS}": address,
        "{COMMISSION_INFO}": commission_info,
        "{CHAIRMAN}": chairman,
        "{MEMBERS}": ", ".join(members),
        "{OWNER}": owner,
        "{DOCUMENTS}": documents,
        "{CONCLUSION}": conclusion,
        "{JUSTIFICATION}": justification,
    }

    def replace_in_paragraph(paragraph):
        text = paragraph.text
        for key, value in replacements.items():
            if key in text:
                text = text.replace(key, value)
        if text != paragraph.text:
            paragraph.clear()
            run = paragraph.add_run(text)
            run.font.underline = True

    # обработка абзацев
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    # обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

    # ==== ДОБАВЛЕНИЕ ПОДПИСЕЙ ВНИЗУ ====
    doc.add_paragraph("   Председатель межведомственной комиссии:")

    p = doc.add_paragraph()
    p.add_run("\t_______________\t\t\t\t________________________\n")
    doc.add_paragraph()
    p.add_run("\t      подпись\t\t\t\t\t\t")
    p.add_run(chairman)

    doc.add_paragraph("\n   Члены межведомственной комиссии:\n")
    for member in members:
        p = doc.add_paragraph()
        p.add_run("\t_______________\t\t\t\t________________________\n")
        doc.add_paragraph()
        p.add_run("\t     подпись\t\t\t\t\t\t")
        p.add_run(member)

    doc.save(output_path)
